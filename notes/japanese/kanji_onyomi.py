import os
import copy

def readfile():
    data = []
    dirname = 'kanji'
    for f in os.listdir(dirname):
        fpath = os.path.join(dirname, f)
        h = open(fpath, 'r')
        for i in h.readlines():
            if i.strip() == '':
                continue
            data.append(i.strip())
        h.close()
    return data


def classify(source):
    kanji_set = {}
    
    kanji = {}
    previous = '*'
    ji = []
    ji_in = {}
    for current in source:
        if current == '':
            continue
        if current == '*' and previous == '*':
            continue
        if current == '*':
            previous = current
            kanji_set['/'.join(ji)] = kanji
            ji = []
            kanji = {}
        else:
            if len(current) == 1:
                ji.append(current)
            else:
                items = [x.strip() for x in current.split('\t')]
                if len(items) > 3:
                    print(items)
                elif len(items) == 3 and len(ji) == 0 and len(items[0]) == 1:
                    ji.append(items[0])
                    kanji[items[1]] = items[2]
                elif len(items) == 3:
                    print(items)
                elif len(items) == 2:
                    kanji[items[0]] = items[1]
                else:
                    kanji[''] = current
            previous = current

    return kanji_set


def reorgnize(kanji_set):
    new_kanjiset = {}
    for a, y in kanji_set.items():
        for b, c in y.items():
            if b not in new_kanjiset:
                new_kanjiset[b] = {}
            new_kanjiset[b][a] = c
    return new_kanjiset

def check_katakana_hirakana(word):
    all_hira = True
    all_kata = True
    for i in word:
        all_hira = all_hira and u'\u3040' <= i <= u'\u309F'
        all_kata = all_kata and u'\u30A0' <= i <= u'\u30FF'
    return all_hira, all_kata

def get_result(kanji):
    hira_set = {}
    kata_set = {}
    pointer = None
    for x, y in kanji.items():
        hira, kata = check_katakana_hirakana(x)
        if hira and kata or not hira and not kata:
            print(x, y)
            continue
        pointer = hira_set if hira else kata_set
        pointer[x] = y
    return hira_set, kata_set

def add_alter(sets):
    tmp = {}
    for i, a in sets.items():
        for j in a.keys():
            if j not in tmp:
                tmp[j] = []
            tmp[j].append(i)
    multi = {x: y for x, y in tmp.items() if len(y) >= 2}
    
    # For kanji which has one sound, or which share sound with other kanji
    newsets = {}
    # For kanji which has 2 ore more sounds and do not share any sound with other kanji
    remains = {}
    multi_saver = {}
    for i, a in sets.items():

        tmp = {}
        for j, k in a.items():
            if j not in multi:
                continue
            for p in multi[j]:
                if p == i:
                    continue
                z = '/'.join(sorted([i,p]))
                tmp[z] = tmp.get(z, [])
                tmp[z].append(j)
            if z in multi_saver and multi_saver[z] != tmp[z]:
                print(multi_saver[z], tmp[z])
            else:
                multi_saver[z] = tmp[z]
        tmp_sounds = []
        tmp_kanji = []
        for x, y in tmp.items():
            if len(y) <= 1:
                continue
            tmp_sounds += x.split('/')
            tmp_kanji += y
            newsets[x] = newsets.get(x, {})
            for j in y:
                newsets[x][j] = a[j]

        for j, k in a.items():
            if j in tmp_kanji:
                continue
            # For kanji which has one sound, or which share sound with other kanji
            l = '%s(%s)' %(j, '/'.join([x for x in multi[j] if i != x])) if j in multi else j
            newsets[i] = newsets.get(i, {})
            newsets[i][l] = k

            # For kanji which has 2 or more sounds and do not share any sound with other kanji
            if j in multi:
                if j not in remains:
                    remains[j] = {}
                remains[j][i] = k

    return newsets, remains


def save_sets(sets, name):
    import docx
    doc = docx.Document()
    doc.add_heading('常用汉字列表%s' %name)
    for a in sorted(sets.keys()):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = docx.shared.Inches(-1)
        p.paragraph_format.right_indent = docx.shared.Inches(-1)
        p.paragraph_format.line_spacing = docx.shared.Inches(0)
        p.paragraph_format.space_before = docx.shared.Inches(0)
        p.paragraph_format.space_after = docx.shared.Inches(0)

        yomi = p.add_run(a + '\t')
        yomi.bold = True
        yomi.font.size = docx.shared.Pt(16)

        kanji = p.add_run('、'.join(sets[a].keys()))
        kanji.bold = False
        kanji.font.size = docx.shared.Pt(10)

    doc.save('/mnt/c/Users/cj/Desktop/test.docx')

def save_sets_old2(sets, name):
    import docx
    doc = docx.Document()
    doc.add_heading('常用汉字列表%s' %name)
    
    p = doc.add_table(rows=1, cols=2)
    p.style.paragraph_format.left_indent = docx.shared.Inches(-1)
    p.rows[0].cells[0].text = name
    p.rows[0].cells[1].text = '漢字'
    for a in sorted(sets.keys()):
        r = p.add_row().cells
        r[0].text = a
        r[1].text = '、'.join(sets[a].keys())

    doc.save('/mnt/c/Users/cj/Desktop/test.docx')


def save_sets_csv(sets, name):
    h = open('/mnt/c/Users/cj/Desktop/常用汉字列表%s' %name, 'w')
    h.write('%s, 漢字\n' %name)
    for a in sorted(sets.keys()):
        h.write(a + ',' + '、'.join(sets[a].keys()) + '\n')
    h.close()



source = readfile()
kanji  = classify(source)
newkanji = reorgnize(kanji)
hira, kata = get_result(newkanji)
kata, remains = add_alter(kata)
save_sets(kata, '音読み')
#save_sets(remains)
