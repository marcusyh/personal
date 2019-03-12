import docx

def save_to_docx(sets, filename, filepath):
    doc = docx.Document()
    doc.add_heading(filename)
    for a in sorted(sets.keys()):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = docx.shared.Inches(-1)
        p.paragraph_format.right_indent = docx.shared.Inches(-1)
        p.paragraph_format.line_spacing = docx.shared.Inches(0)
        p.paragraph_format.space_before = docx.shared.Inches(0)
        p.paragraph_format.space_after = docx.shared.Inches(0)

        yomi = p.add_run(a + '\t')
        yomi.bold = True
        yomi.font.size = docx.shared.Pt(16)

        kanji = p.add_run('、'.join(sets[a].keys()))
        kanji.bold = False
        kanji.font.size = docx.shared.Pt(10)

    doc.save('%s/%s.docx' %(filepath, filename))

def save_to_docx_old(sets, filename, filepath):
    import docx
    doc = docx.Document()
    doc.add_heading(filename)
    
    p = doc.add_table(rows=1, cols=2)
    p.style.paragraph_format.left_indent = docx.shared.Inches(-1)
    p.rows[0].cells[0].text = name
    p.rows[0].cells[1].text = '漢字'
    for a in sorted(sets.keys()):
        r = p.add_row().cells
        r[0].text = a
        r[1].text = '、'.join(sets[a].keys())

    doc.save('%s/%s.docx' %(filepath, filename))


def save_to_csv(sets, filename, filepath):
    h = open('%s/%s.csv' %(filepath, filename), 'w')
    h.write('%s, 漢字\n' %filename)
    for a in sorted(sets.keys()):
        h.write(a + ',' + '、'.join(sets[a].keys()) + '\n')
    h.close()

